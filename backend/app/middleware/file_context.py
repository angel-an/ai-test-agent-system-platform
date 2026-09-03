"""
File Context Middleware - Multi-file context extraction for conversations.

Supports: PDF, images (png/jpg/gif/webp), Excel (.xlsx/.xls), Word (.docx/.doc),
text files (.txt/.md/.csv/.json/.xml/.yaml/.yml/.log)

Processing strategy:
- PDF: Reuse existing pdf processor (PyMuPDF4LLM)
- Images: Doubao multimodal vision model for OCR/description
- Excel: openpyxl for .xlsx, xlrd for .xls
- Word: python-docx for .docx; pywin32 COM for .doc (convert to .docx then parse)
- Text: Direct read with encoding detection

Adapted from the reference project to work with the current platform's
configuration system (settings.py).
"""

from __future__ import annotations

import asyncio
import base64
import hashlib
import logging
import os
import tempfile
import time
from collections import OrderedDict
from pathlib import Path
from typing import Any, Awaitable, Callable

from langchain.agents.middleware import AgentMiddleware, ModelRequest, ModelResponse
from langchain.agents.middleware.types import ResponseT
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_openai import ChatOpenAI
from langgraph.typing import ContextT

from app.config.settings import settings

logger = logging.getLogger(__name__)


class _SimpleLRUCache:
    """Simple LRU cache using OrderedDict (no external dependencies)."""

    def __init__(self, maxsize: int = 100):
        self.maxsize = maxsize
        self._cache: OrderedDict[str, Any] = OrderedDict()

    def __contains__(self, key: str) -> bool:
        return key in self._cache

    def __getitem__(self, key: str) -> Any:
        value = self._cache.pop(key)
        self._cache[key] = value
        return value

    def __setitem__(self, key: str, value: Any) -> None:
        if key in self._cache:
            self._cache.pop(key)
        self._cache[key] = value
        if len(self._cache) > self.maxsize:
            self._cache.popitem(last=False)

    def get(self, key: str, default: Any = None) -> Any:
        if key in self._cache:
            return self[key]
        return default

    def __len__(self) -> int:
        return len(self._cache)


# Upload directory — backend/workspace/uploads
_UPLOAD_DIR = (
    Path(__file__).resolve().parent.parent.parent / "workspace" / "uploads"
)
_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

# File content cache (LRU: max 100 entries)
_file_cache = _SimpleLRUCache(maxsize=100)

# Supported MIME types
SUPPORTED_MIME_TYPES = {
    # PDF
    "application/pdf",
    # Images
    "image/png",
    "image/jpeg",
    "image/gif",
    "image/webp",
    # Excel
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",  # .xlsx
    "application/vnd.ms-excel",  # .xls
    # Word
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # .docx
    "application/msword",  # .doc
    # Text
    "text/plain",
    "text/markdown",
    "text/csv",
    "application/json",
    "application/xml",
    "text/xml",
    "text/yaml",
    "application/x-yaml",
}

# Extension to MIME type fallback mapping
EXTENSION_MIME_MAP = {
    ".pdf": "application/pdf",
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".gif": "image/gif",
    ".webp": "image/webp",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".xls": "application/vnd.ms-excel",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".doc": "application/msword",
    ".txt": "text/plain",
    ".md": "text/markdown",
    ".csv": "text/csv",
    ".json": "application/json",
    ".xml": "application/xml",
    ".yaml": "text/yaml",
    ".yml": "text/yaml",
    ".log": "text/plain",
}


def _decode_base64(data: str) -> bytes:
    """Decode base64 string to bytes, stripping data URI prefix if present."""
    if "," in data:
        data = data.split(",", 1)[1]
    return base64.b64decode(data)


def _get_cache_key(data: bytes, filename: str) -> str:
    """Generate cache key from file content hash and filename."""
    file_hash = hashlib.md5(data).hexdigest()
    return f"{filename}_{file_hash}"


def _get_mime_type(filename: str, declared_mime: str = "") -> str:
    """Determine MIME type from declared type or file extension."""
    if declared_mime and declared_mime in SUPPORTED_MIME_TYPES:
        return declared_mime
    ext = Path(filename).suffix.lower()
    return EXTENSION_MIME_MAP.get(ext, "")


# ---------------------------------------------------------------------------
# File Processors
# ---------------------------------------------------------------------------


def extract_image_text(image_data: bytes, filename: str, mime_type: str) -> str:
    """
    Use Doubao multimodal vision model to extract text from an image.

    Args:
        image_data: Raw image bytes
        filename: Original filename
        mime_type: Image MIME type

    Returns:
        Extracted text description from the image
    """
    cache_key = _get_cache_key(image_data, filename)
    if cache_key in _file_cache:
        logger.info("Image cache hit: %s", filename)
        return _file_cache[cache_key]

    try:
        api_key = settings.image_parser_api_key
        if not api_key:
            logger.warning("image_parser_api_key not configured, skipping image analysis")
            return "[图片解析跳过：未配置多模态API Key]"

        model = ChatOpenAI(
            base_url=settings.image_parser_api_base,
            api_key=api_key,
            model=settings.image_parser_model,
        )

        b64_data = base64.b64encode(image_data).decode("utf-8")
        image_url = f"data:{mime_type};base64,{b64_data}"

        prompt = (
            "请详细描述这张图片的内容。如果是界面截图，请提取所有可见的文字、"
            "按钮、输入框、列表等元素，并描述页面布局和交互逻辑。"
            "如果是流程图或架构图，请描述其中的节点、连线和关系。"
        )

        message = HumanMessage(
            content=[
                {"type": "text", "text": prompt},
                {"type": "image_url", "image_url": {"url": image_url}},
            ]
        )

        response = model.invoke([message])
        text_content = (
            response.content
            if isinstance(response.content, str)
            else str(response.content)
        )

        logger.info(
            "Image parsed successfully: %s, length: %d chars",
            filename,
            len(text_content),
        )
        _file_cache[cache_key] = text_content
        return text_content

    except Exception as e:
        logger.error("Image parsing failed for %s: %s", filename, e)
        return f"[图片解析失败: {str(e)}]"


def extract_excel_text(excel_data: bytes, filename: str) -> str:
    """
    Extract text content from Excel files (.xlsx / .xls).

    Args:
        excel_data: Raw Excel bytes
        filename: Original filename

    Returns:
        Extracted text in Markdown table format
    """
    cache_key = _get_cache_key(excel_data, filename)
    if cache_key in _file_cache:
        logger.info("Excel cache hit: %s", filename)
        return _file_cache[cache_key]

    ext = Path(filename).suffix.lower()
    temp_path = None

    try:
        temp_file = tempfile.NamedTemporaryFile(suffix=ext, delete=False)
        temp_file.write(excel_data)
        temp_file.flush()
        os.fsync(temp_file.fileno())
        temp_path = temp_file.name
        temp_file.close()

        if ext == ".xlsx":
            text_content = _parse_xlsx(temp_path)
        elif ext == ".xls":
            text_content = _parse_xls(temp_path)
        else:
            text_content = f"Unsupported Excel format: {ext}"

        logger.info(
            "Excel parsed successfully: %s, length: %d chars", filename, len(text_content)
        )
        _file_cache[cache_key] = text_content
        return text_content

    except Exception as e:
        logger.error("Excel parsing failed for %s: %s", filename, e)
        return f"[Excel解析失败: {str(e)}]"
    finally:
        if temp_path and os.path.exists(temp_path):
            _safe_delete(temp_path)


def _parse_xlsx(file_path: str) -> str:
    """Parse .xlsx file using openpyxl."""
    from openpyxl import load_workbook

    wb = load_workbook(file_path, read_only=True, data_only=True)
    result_parts = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue

        result_parts.append(f"## Sheet: {sheet_name}\n")

        for row_idx, row in enumerate(rows):
            cells = [str(cell) if cell is not None else "" for cell in row]
            result_parts.append("| " + " | ".join(cells) + " |")
            if row_idx == 0:
                result_parts.append("| " + " | ".join(["---"] * len(cells)) + " |")

        result_parts.append("")

    wb.close()
    return "\n".join(result_parts) if result_parts else "Excel file is empty"


def _parse_xls(file_path: str) -> str:
    """Parse .xls file using xlrd."""
    try:
        import xlrd
    except ImportError:
        return "xlrd library not installed, cannot parse .xls files. Install with: pip install xlrd"

    wb = xlrd.open_workbook(file_path)
    result_parts = []

    for sheet_idx in range(wb.nsheets):
        ws = wb.sheet_by_index(sheet_idx)
        if ws.nrows == 0:
            continue

        result_parts.append(f"## Sheet: {ws.name}\n")

        for row_idx in range(ws.nrows):
            cells = [
                str(ws.cell_value(row_idx, col_idx)) for col_idx in range(ws.ncols)
            ]
            result_parts.append("| " + " | ".join(cells) + " |")
            if row_idx == 0:
                result_parts.append("| " + " | ".join(["---"] * len(cells)) + " |")

        result_parts.append("")

    return "\n".join(result_parts) if result_parts else "Excel file is empty"


def _doc_to_docx(doc_path: str) -> str | None:
    """Convert old .doc to .docx (Windows COM interface).

    Supports Microsoft Word and WPS Office (via kwps.Application).

    Args:
        doc_path: .doc file path

    Returns:
        Path to converted .docx temp file, or None on failure
    """
    temp_docx = doc_path + ".converted.docx"

    # Try Microsoft Word
    try:
        import win32com.client as wc

        word = wc.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(doc_path)
        doc.SaveAs(temp_docx, 16)  # 16 = wdFormatXMLDocument (.docx)
        doc.Close()
        word.Quit()
        logger.info("[doc→docx] Word conversion success: %s → %s", doc_path, temp_docx)
        return temp_docx
    except Exception as e_word:
        logger.debug("[doc→docx] Word conversion failed: %s", e_word)

    # Try WPS Office
    try:
        import win32com.client as wc

        wps = wc.Dispatch("kwps.Application")
        wps.Visible = False
        doc = wps.Documents.Open(doc_path)
        try:
            doc.SaveAs2(temp_docx, 16)
        except Exception:
            doc.SaveAs(temp_docx, 16)
        doc.Close()
        wps.Quit()
        logger.info("[doc→docx] WPS conversion success: %s → %s", doc_path, temp_docx)
        return temp_docx
    except Exception as e_wps:
        logger.debug("[doc→docx] WPS conversion failed: %s", e_wps)

    logger.warning(
        "[doc→docx] Conversion failed (Word or WPS not found): %s", doc_path
    )
    return None


def _extract_docx_images(doc) -> list[tuple[bytes, str]]:
    """Extract all embedded images from a docx Document (deduplicated).

    Args:
        doc: python-docx Document object

    Returns:
        List of (image_bytes, mime_type) tuples
    """
    images = []
    seen_hashes: set[str] = set()

    for rel in doc.part.rels.values():
        if "image" not in getattr(rel, "reltype", ""):
            continue
        try:
            image_part = rel.target_part
            image_data = image_part.blob
        except Exception:
            continue

        # Deduplicate: same image may be referenced multiple times
        h = hashlib.md5(image_data).hexdigest()
        if h in seen_hashes:
            continue
        seen_hashes.add(h)

        content_type = getattr(image_part, "content_type", "")
        if "png" in content_type:
            mime = "image/png"
        elif "jpeg" in content_type or "jpg" in content_type:
            mime = "image/jpeg"
        elif "gif" in content_type:
            mime = "image/gif"
        elif "bmp" in content_type:
            mime = "image/bmp"
        elif "webp" in content_type:
            mime = "image/webp"
        else:
            mime = "image/png"  # fallback

        images.append((image_data, mime))

    return images


def extract_word_text(word_data: bytes, filename: str) -> str:
    """
    Extract text content from Word documents (.docx / .doc).

    Also extracts and analyzes embedded images using the vision model.

    Args:
        word_data: Raw Word document bytes
        filename: Original filename

    Returns:
        Extracted text content with image descriptions
    """
    cache_key = _get_cache_key(word_data, filename)
    if cache_key in _file_cache:
        logger.info("Word cache hit: %s", filename)
        return _file_cache[cache_key]

    temp_path = None
    converted_docx = None
    try:
        ext = Path(filename).suffix.lower()
        temp_file = tempfile.NamedTemporaryFile(suffix=ext, delete=False)
        temp_file.write(word_data)
        temp_file.flush()
        os.fsync(temp_file.fileno())
        temp_path = temp_file.name
        temp_file.close()

        # Convert old .doc to .docx
        if ext == ".doc":
            converted_docx = _doc_to_docx(temp_path)
            if converted_docx:
                doc_path = converted_docx
            else:
                return "旧版 .doc 文件转换失败（需 Windows + Office/Word）。建议将文档另存为 .docx 后重新上传。"
        else:
            doc_path = temp_path

        from docx import Document

        doc = Document(doc_path)
        parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if para.style and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "").replace("Heading", "1")
                try:
                    level = int(level)
                except ValueError:
                    level = 1
                parts.append(f"{'#' * level} {text}")
            else:
                parts.append(text)

        # Extract tables
        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                parts.append("| " + " | ".join(cells) + " |")
                if row_idx == 0:
                    parts.append("| " + " | ".join(["---"] * len(cells)) + " |")
            parts.append("")

        text_content = "\n".join(parts) if parts else "Word document is empty"

        # Extract and analyze embedded images
        try:
            images = _extract_docx_images(doc)
            if images:
                logger.info(
                    "[extract_word_text] Found %d embedded images: %s",
                    len(images),
                    filename,
                )
                image_parts = []
                for idx, (img_data, mime) in enumerate(images[:20], 1):  # limit 20
                    try:
                        desc = extract_image_text(
                            img_data, f"{filename}_img_{idx}", mime
                        )
                        image_parts.append(
                            f"[嵌入图片 {idx}]\n{desc}\n[/嵌入图片 {idx}]"
                        )
                    except Exception as img_e:
                        logger.warning(
                            "[extract_word_text] Image %d parse failed: %s", idx, img_e
                        )
                        image_parts.append(
                            f"[嵌入图片 {idx}]\n[图片解析失败: {img_e}]\n[/嵌入图片 {idx}]"
                        )
                if image_parts:
                    text_content += (
                        "\n\n## 文档内嵌图片解析\n\n" + "\n\n".join(image_parts)
                    )
        except Exception as e:
            logger.warning("[extract_word_text] Image extraction failed: %s", e)

        logger.info(
            "Word parsed successfully: %s, length: %d chars", filename, len(text_content)
        )
        _file_cache[cache_key] = text_content
        return text_content

    except ImportError:
        return "python-docx library not installed. Install with: pip install python-docx"
    except Exception as e:
        logger.error("Word parsing failed for %s: %s", filename, e)
        return f"[Word解析失败: {str(e)}]"
    finally:
        if temp_path and os.path.exists(temp_path):
            _safe_delete(temp_path)
        if converted_docx and os.path.exists(converted_docx):
            _safe_delete(converted_docx)


def extract_text_file(text_data: bytes, filename: str) -> str:
    """
    Extract content from text-based files.

    Args:
        text_data: Raw file bytes
        filename: Original filename

    Returns:
        File text content
    """
    cache_key = _get_cache_key(text_data, filename)
    if cache_key in _file_cache:
        logger.info("Text file cache hit: %s", filename)
        return _file_cache[cache_key]

    try:
        for encoding in ("utf-8", "gbk", "gb2312", "latin-1"):
            try:
                text_content = text_data.decode(encoding)
                break
            except (UnicodeDecodeError, LookupError):
                continue
        else:
            text_content = text_data.decode("utf-8", errors="replace")

        logger.info(
            "Text file read successfully: %s, length: %d chars",
            filename,
            len(text_content),
        )
        _file_cache[cache_key] = text_content
        return text_content

    except Exception as e:
        logger.error("Text file reading failed for %s: %s", filename, e)
        return f"[文本文件读取失败: {str(e)}]"


def _safe_delete(file_path: str, max_retries: int = 3, delay: float = 0.1) -> None:
    """Safely delete temporary file with retry for Windows file locking."""
    if not os.path.exists(file_path):
        return
    for attempt in range(max_retries):
        try:
            os.unlink(file_path)
            return
        except PermissionError:
            if attempt < max_retries - 1:
                time.sleep(delay)
            else:
                logger.warning(
                    "Cannot delete temp file after %d retries: %s", max_retries, file_path
                )
        except Exception as e:
            logger.warning("Error deleting temp file: %s", e)
            break


# ---------------------------------------------------------------------------
# PDF integration with existing processor
# ---------------------------------------------------------------------------


def extract_pdf_text(pdf_data: bytes, filename: str) -> str:
    """Extract text from PDF using the existing pdf_processor module."""
    try:
        from app.agents.tools.testcase.pdf_processor import extract_pdf_text as _extract_pdf

        return _extract_pdf(pdf_data, filename)
    except Exception as e:
        logger.error("PDF extraction failed for %s: %s", filename, e)
        return f"[PDF解析失败: {str(e)}]"


# ---------------------------------------------------------------------------
# Main file content router
# ---------------------------------------------------------------------------


def extract_file_content(file_data: bytes, filename: str, mime_type: str = "") -> str:
    """
    Route file to appropriate processor based on MIME type.

    Args:
        file_data: Raw file bytes
        filename: Original filename
        mime_type: MIME type (auto-detected from extension if empty)

    Returns:
        Extracted text content
    """
    resolved_mime = _get_mime_type(filename, mime_type)

    if not resolved_mime:
        return f"[不支持的文件类型: {filename}]"

    if resolved_mime == "application/pdf":
        return extract_pdf_text(file_data, filename)

    if resolved_mime.startswith("image/"):
        return extract_image_text(file_data, filename, resolved_mime)

    if resolved_mime in (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    ):
        return extract_excel_text(file_data, filename)

    if resolved_mime in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        return extract_word_text(file_data, filename)

    if resolved_mime.startswith("text/") or resolved_mime in (
        "application/json",
        "application/xml",
        "application/x-yaml",
    ):
        return extract_text_file(file_data, filename)

    return f"[不支持的文件类型: {resolved_mime}]"


# ---------------------------------------------------------------------------
# Middleware
# ---------------------------------------------------------------------------

# Maximum characters per file to inject into LLM context.
_MAX_FILE_CHARS = 150_000

# Maximum total context chars from all uploaded files combined.
_MAX_TOTAL_CONTEXT_CHARS = 400_000

# Process-level context dedup cache: {thread_id: context_hash}
_global_context_hashes: dict[str, str] = {}


class FileContextMiddleware(AgentMiddleware):
    """Multi-file context extraction middleware.

    Persists every uploaded file's extracted text across the entire conversation
    so the agent can answer questions about a document several turns later.

    Strategy:
      - Track every attachment seen in any HumanMessage by content hash.
      - On each model call, append a single SystemMessage containing the merged
        text of all files known so far. The original messages are not mutated.
      - Large files are truncated to prevent exceeding the LLM's context limit.
      - Process-level dedup: if the same context was already injected for this
        thread, skip re-injection to save bandwidth and tokens.
    """

    def __init__(self) -> None:
        super().__init__()
        self._files_seen: dict[str, tuple[str, str]] = {}

    def _get_thread_id(self, request: ModelRequest) -> str:
        """Extract thread_id from request for process-level cache."""
        try:
            runtime = getattr(request, "runtime", None)
            if runtime is not None:
                cfg = getattr(runtime, "config", None) or {}
                if isinstance(cfg, dict):
                    configurable = cfg.get("configurable")
                    if isinstance(configurable, dict):
                        tid = configurable.get("thread_id")
                        if tid:
                            return str(tid)
        except Exception:
            pass

        # Fallback: use last message id as weak identifier
        try:
            msgs = getattr(request, "messages", [])
            if msgs and len(msgs) > 0:
                last_msg = msgs[-1]
                if hasattr(last_msg, "id") and last_msg.id:
                    return str(last_msg.id)
        except Exception:
            pass

        return "_unknown_thread"

    async def awrap_model_call(
        self,
        request: ModelRequest[ContextT],
        handler: Callable[[ModelRequest[ContextT]], Awaitable[ModelResponse[ResponseT]]],
    ) -> Any:
        if not request.messages:
            return await handler(request)

        # 1. Scan every HumanMessage for new attachments and parse them.
        new_files: list[tuple[bytes, str, str]] = []
        for msg in request.messages:
            if not isinstance(msg, HumanMessage):
                continue
            for file_data, filename, mime_type in self._extract_files_from_message(msg):
                cache_key = _get_cache_key(file_data, filename)
                if cache_key in self._files_seen:
                    continue
                self._files_seen[cache_key] = (filename, "")
                new_files.append((file_data, filename, mime_type))

        if new_files:
            # Save originals to uploads dir so tools can access by path.
            for file_data, filename, _ in new_files:
                safe_name = Path(filename).name
                file_path = _UPLOAD_DIR / safe_name
                try:
                    file_path.write_bytes(file_data)
                    logger.info("[FileContextMiddleware] File saved: %s", file_path)
                except Exception as e:
                    logger.warning(
                        "[FileContextMiddleware] Failed to save %s: %s", filename, e
                    )

            # Heavy I/O off the event loop.
            for file_data, filename, mime_type in new_files:
                cache_key = _get_cache_key(file_data, filename)
                try:
                    text = await asyncio.to_thread(
                        extract_file_content, file_data, filename, mime_type
                    )
                except Exception as e:
                    text = f"[文件解析失败: {e}]"
                self._files_seen[cache_key] = (filename, text)

        # 2. If no files have ever been uploaded, pass through.
        if not self._files_seen:
            return await handler(request)

        # 3. Build a single SystemMessage with all known file contents.
        sections = []
        total_chars = 0
        for idx, (filename, text) in enumerate(self._files_seen.values(), 1):
            if not text:
                continue

            original_len = len(text)
            if original_len > _MAX_FILE_CHARS:
                text = text[:_MAX_FILE_CHARS] + (
                    f"\n\n[Content truncated: {original_len:,} chars total, "
                    f"only first {_MAX_FILE_CHARS:,} chars shown. "
                    f"Use tools to read specific sections if needed.]"
                )

            section = f"--- File {idx}: {filename} ---\n{text}"
            total_chars += len(section)
            sections.append(section)

        if not sections:
            return await handler(request)

        file_names = [info[0] for info in self._files_seen.values()]
        header = (
            f"[Uploaded Documents Context]\n"
            f"The user has uploaded {len(file_names)} document(s) in this conversation: "
            f"{', '.join(file_names)}.\n"
            f"Treat the content below as the authoritative requirement source. "
            f"When the user references the upload (\"已上传的需求\", a filename, etc.), "
            f"use this content directly without asking them to re-upload.\n\n"
        )

        body = "\n\n".join(sections)
        system_text = header + body

        # Final safety truncation
        if len(system_text) > _MAX_TOTAL_CONTEXT_CHARS:
            allowed_body = _MAX_TOTAL_CONTEXT_CHARS - len(header) - 200
            body = body[:allowed_body] + (
                f"\n\n[Total uploaded context truncated: "
                f"{len(system_text):,} chars total, only {_MAX_TOTAL_CONTEXT_CHARS:,} chars shown.]"
            )
            system_text = header + body
            logger.warning(
                "[FileContextMiddleware] Total context truncated to %d chars",
                _MAX_TOTAL_CONTEXT_CHARS,
            )

        # Truncation awareness warning
        has_truncation = (
            "[Content truncated:" in body
            or "[Total uploaded context truncated:" in body
        )
        if has_truncation:
            truncation_warning = (
                "[⚠️ 文档截断警告] 上传的文档超出注入上限（单文件 150K 字符 / 总计 400K 字符），"
                "部分内容未被加载到上下文中。\n"
                "请在 Step 1a 文档结构确认中说明此情况，并建议用户：\n"
                "1. 拆分文档分多次上传，或\n"
                "2. 仅上传与当前分析相关的章节。\n\n"
            )
            system_text = truncation_warning + system_text
            logger.warning(
                "[FileContextMiddleware] Truncation warning injected for thread %s",
                self._get_thread_id(request),
            )

        # Process-level dedup
        thread_id = self._get_thread_id(request)
        context_hash = hashlib.md5(system_text.encode()).hexdigest()
        global _global_context_hashes
        if _global_context_hashes.get(thread_id) == context_hash:
            logger.debug(
                "[FileContextMiddleware] Context unchanged for thread %s, skipping",
                thread_id,
            )
            return await handler(request)
        _global_context_hashes[thread_id] = context_hash
        logger.info(
            "[FileContextMiddleware] Injecting context for thread %s (hash=%s... len=%d chars)",
            thread_id,
            context_hash[:8],
            len(system_text),
        )

        new_messages = list(request.messages) + [SystemMessage(content=system_text)]
        request = request.override(messages=new_messages)

        return await handler(request)

    def _extract_files_from_message(
        self, msg: HumanMessage
    ) -> list[tuple[bytes, str, str]]:
        """
        Extract all supported file attachments from the message.

        Returns:
            List of (file_bytes, filename, mime_type) tuples
        """
        attachments = msg.additional_kwargs.get("attachments", [])
        if not isinstance(attachments, list):
            return []

        files = []
        for att in attachments:
            if not isinstance(att, dict):
                continue

            mime_type = att.get("mimeType", "")
            data = att.get("data")
            if not data or not isinstance(data, str):
                continue

            metadata = att.get("metadata", {})
            filename = metadata.get("filename") or metadata.get("name", "unknown")

            resolved_mime = _get_mime_type(filename, mime_type)
            if not resolved_mime:
                logger.warning(
                    "[FileContextMiddleware] Unsupported file: %s (%s)",
                    filename,
                    mime_type,
                )
                continue

            try:
                file_bytes = _decode_base64(data)
                files.append((file_bytes, filename, resolved_mime))
                logger.info(
                    "[FileContextMiddleware] Extracted: %s (%s)", filename, resolved_mime
                )
            except Exception as e:
                logger.warning(
                    "[FileContextMiddleware] Failed to decode %s: %s", filename, e
                )
                continue

        return files
