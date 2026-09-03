"""API document parser service.

Converts lightweight TXT/Word API documents into an OpenAPI-like spec that can
reuse the existing OpenAPIParser import flow.
"""

from __future__ import annotations

import io
import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


HTTP_METHODS = {"GET", "POST", "PUT", "DELETE", "PATCH", "HEAD", "OPTIONS"}


@dataclass
class APIDocEndpointDraft:
    """Intermediate endpoint extracted from a TXT/Word API document."""

    method: str
    path: str
    summary: str | None = None
    tag: str = "API文档"
    description: str | None = None
    request_example: Any | None = None
    response_example: Any | None = None
    extra_response_fields: dict[str, dict[str, Any]] = field(default_factory=dict)


class APIDocParserService:
    """Parse non-OpenAPI API docs and convert them into OpenAPI specs."""

    endpoint_pattern = re.compile(
        r"^\s*(GET|POST|PUT|DELETE|PATCH|HEAD|OPTIONS)\s+(/[^\s]+)",
        re.IGNORECASE,
    )
    base_url_pattern = re.compile(
        r"(?:域名|base\s*url|baseUrl|服务器地址|服务地址)\s*[:：]\s*(https?://[^\s]+)",
        re.IGNORECASE,
    )
    heading_pattern = re.compile(r"^\s*--\s+(.+?)\s*$")
    extra_response_field_pattern = re.compile(
        r"(?:增加|新增)?返回参数\s+([A-Za-z_][A-Za-z0-9_]*)\s*(.*)",
        re.IGNORECASE,
    )
    curl_command_pattern = re.compile(r"(?:^|\n)\s*curl\b(?P<command>.*?)(?=\n\s*\n|\Z)", re.IGNORECASE | re.DOTALL)
    curl_method_pattern = re.compile(r"(?:--request|-X)\s+['\"]?(GET|POST|PUT|DELETE|PATCH|HEAD|OPTIONS)['\"]?", re.IGNORECASE)
    curl_url_pattern = re.compile(r"['\"](https?://[^'\"\s]+)['\"]", re.IGNORECASE)
    curl_data_pattern = re.compile(r"(?:--data-raw|--data|-d)\s+(['\"])(?P<data>.*?)\1", re.IGNORECASE | re.DOTALL)
    sensitive_request_field_pattern = re.compile(r"(?:phone|mobile|tel|telephone|password|token|authorization)", re.IGNORECASE)

    def parse_text_to_draft(self, text: str, title: str | None = None) -> dict[str, Any]:
        """Parse text content into a normalized draft model."""

        normalized = text.replace("\r\n", "\n").replace("\r", "\n")
        lines = normalized.split("\n")
        base_url = self._extract_base_url(normalized)
        endpoint_positions = self._find_endpoint_positions(lines)

        endpoints: list[APIDocEndpointDraft] = []
        for index, line_no in enumerate(endpoint_positions):
            match = self.endpoint_pattern.match(lines[line_no])
            if not match:
                continue

            next_line_no = (
                endpoint_positions[index + 1]
                if index + 1 < len(endpoint_positions)
                else len(lines)
            )
            headings = self._collect_headings_before(lines, line_no)
            summary = headings[-1] if headings else None
            tag = headings[0] if len(headings) > 1 else "API文档"
            segment_lines = lines[line_no + 1 : next_line_no]

            endpoint = APIDocEndpointDraft(
                method=match.group(1).upper(),
                path=match.group(2).strip(),
                summary=summary,
                tag=tag,
                description=self._build_description(segment_lines),
                request_example=self._extract_marked_json(segment_lines, ["请求体", "请求参数", "Request Body"]),
                response_example=self._extract_marked_json(segment_lines, ["响应", "返回", "Response"]),
                extra_response_fields=self._extract_extra_response_fields(segment_lines),
            )
            endpoints.append(endpoint)

        # curl is common in delivery documents. Parse it separately because a
        # command spans multiple lines and contains headers that must not enter
        # generated API assets.
        endpoints.extend(self._parse_curl_endpoints(normalized))
        if not base_url:
            base_url = self._infer_base_url_from_curl(normalized)

        return {
            "title": title or self._infer_title(lines),
            "base_url": base_url,
            "endpoints": [self._endpoint_to_dict(endpoint) for endpoint in endpoints],
            "total_endpoints": len(endpoints),
        }

    def draft_to_openapi(self, draft: dict[str, Any]) -> dict[str, Any]:
        """Convert a normalized draft model into an OpenAPI 3.0 spec."""

        spec: dict[str, Any] = {
            "openapi": "3.0.0",
            "info": {
                "title": draft.get("title") or "API文档",
                "version": "1.0.0",
            },
            "paths": {},
        }
        if draft.get("base_url"):
            spec["servers"] = [{"url": draft["base_url"]}]

        for endpoint in draft.get("endpoints", []):
            path = endpoint["path"]
            method = endpoint["method"].lower()
            operation: dict[str, Any] = {
                "tags": [endpoint.get("tag") or "API文档"],
                "summary": endpoint.get("summary") or f"{endpoint['method']} {path}",
                "description": endpoint.get("description"),
                "responses": self._build_openapi_responses(endpoint),
            }
            if endpoint.get("request_example") is not None:
                request_schema = self._infer_schema(endpoint["request_example"], required_for_object=True)
                operation["requestBody"] = {
                    "required": True,
                    "content": {
                        "application/json": {
                            "schema": request_schema,
                            "example": endpoint["request_example"],
                        }
                    },
                }

            spec["paths"].setdefault(path, {})[method] = operation

        return spec

    def extract_text_from_file(
        self,
        filename: str | None,
        content_type: str | None,
        content: bytes,
    ) -> str:
        """Extract text from TXT/DOCX/DOC content."""

        suffix = Path(filename or "").suffix.lower()
        normalized_type = (content_type or "").lower()

        if suffix in {".txt", ".md", ".text"} or normalized_type.startswith("text/"):
            return self._decode_text(content)

        if suffix == ".docx" or "openxmlformats-officedocument.wordprocessingml.document" in normalized_type:
            try:
                from docx import Document
            except ImportError as exc:
                raise ValueError("解析 .docx 需要安装 python-docx") from exc

            document = Document(io.BytesIO(content))
            parts: list[str] = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        parts.append("\t".join(cells))
            return "\n".join(parts)

        if suffix == ".doc" or normalized_type == "application/msword":
            raise ValueError("暂不直接解析旧版 .doc，请先另存为 .docx 或 txt 后上传")

        raise ValueError(f"不支持的 API 文档类型: {content_type or suffix or 'unknown'}")

    def _extract_base_url(self, text: str) -> str | None:
        match = self.base_url_pattern.search(text)
        return match.group(1).strip() if match else None

    def _parse_curl_endpoints(self, text: str) -> list[APIDocEndpointDraft]:
        endpoints: list[APIDocEndpointDraft] = []
        for match in self.curl_command_pattern.finditer(text):
            command = match.group("command")
            method_match = self.curl_method_pattern.search(command)
            url_match = self.curl_url_pattern.search(command)
            if not method_match or not url_match:
                continue
            method = method_match.group(1).upper()
            url = url_match.group(1)
            path = self._path_from_url(url)
            if not path:
                continue
            data_match = self.curl_data_pattern.search(command)
            request_example = None
            if data_match:
                try:
                    request_example = json.loads(data_match.group("data"))
                except json.JSONDecodeError:
                    request_example = None
            endpoints.append(APIDocEndpointDraft(
                method=method,
                path=path,
                summary=f"{method} {path}",
                tag="API文档",
                description="从 curl 文档解析；授权头不会被导入。",
                request_example=self._redact_request_example(request_example),
            ))
        return endpoints

    def _infer_base_url_from_curl(self, text: str) -> str | None:
        for match in self.curl_command_pattern.finditer(text):
            url_match = self.curl_url_pattern.search(match.group("command"))
            if not url_match:
                continue
            url = url_match.group(1)
            marker = url.find("/api/")
            if marker > 0:
                return url[:marker]
        return None

    @staticmethod
    def _path_from_url(url: str) -> str | None:
        marker = url.find("/api/")
        return url[marker:] if marker >= 0 else None

    def _redact_request_example(self, value: Any, key: str | None = None) -> Any:
        if isinstance(value, dict):
            return {
                item_key: self._redact_request_example(item_value, item_key)
                for item_key, item_value in value.items()
            }
        if isinstance(value, list):
            if key and self.sensitive_request_field_pattern.search(key):
                return [self._redact_request_example(value[0], key)] if value else []
            return [self._redact_request_example(item, key) for item in value]
        if key and self.sensitive_request_field_pattern.search(key):
            return "{{memberPhone}}" if "phone" in key.lower() else "{{runtimeSecret}}"
        return value

    def _find_endpoint_positions(self, lines: list[str]) -> list[int]:
        positions: list[int] = []
        for index, line in enumerate(lines):
            match = self.endpoint_pattern.match(line)
            if match and match.group(1).upper() in HTTP_METHODS:
                positions.append(index)
        return positions

    def _collect_headings_before(self, lines: list[str], endpoint_line_no: int) -> list[str]:
        headings: list[str] = []
        index = endpoint_line_no - 1
        while index >= 0:
            stripped = lines[index].strip()
            if not stripped:
                index -= 1
                continue
            match = self.heading_pattern.match(stripped)
            if match:
                headings.insert(0, match.group(1).strip())
                index -= 1
                continue
            if self.endpoint_pattern.match(stripped):
                break
            if headings:
                break
            index -= 1
        return headings

    def _infer_title(self, lines: list[str]) -> str:
        for line in lines:
            match = self.heading_pattern.match(line)
            if match:
                return match.group(1).strip()
        return "API文档"

    def _build_description(self, segment_lines: list[str]) -> str | None:
        description_lines: list[str] = []
        skip_marked_block = False
        for line in segment_lines:
            stripped = line.strip()
            if not stripped:
                continue
            if self._is_marker_line(stripped):
                skip_marked_block = True
                continue
            if skip_marked_block:
                if stripped.startswith("{") or stripped.startswith("[") or stripped[:1] in {'"', "}", "]", ","}:
                    continue
                skip_marked_block = False
            if self.heading_pattern.match(stripped):
                continue
            if self.base_url_pattern.search(stripped):
                continue
            description_lines.append(stripped)
        return "\n".join(description_lines) or None

    def _is_marker_line(self, line: str) -> bool:
        marker = line.lower()
        return any(token.lower() in marker for token in ["请求体", "请求参数", "request body", "响应", "返回", "response"])

    def _extract_marked_json(self, segment_lines: list[str], markers: list[str]) -> Any | None:
        for index, line in enumerate(segment_lines):
            if not any(marker.lower() in line.lower() for marker in markers):
                continue

            candidates = segment_lines[index + 1 :]
            suffix = re.split(r"[:：]", line, maxsplit=1)
            if len(suffix) == 2 and suffix[1].strip():
                candidates = [suffix[1], *candidates]

            block = self._collect_json_block(candidates)
            if not block:
                return None
            return self._loads_json_with_comments(block)
        return None

    def _collect_json_block(self, lines: list[str]) -> str | None:
        collected: list[str] = []
        started = False
        balance = 0
        in_string = False
        escaped = False

        for raw_line in lines:
            line = raw_line.strip()
            if not started:
                if not line:
                    continue
                first_json_pos = self._find_first_json_char(line)
                if first_json_pos < 0:
                    break
                line = line[first_json_pos:]
                started = True

            collected.append(raw_line)
            for char in self._strip_json_line_comment(raw_line):
                if escaped:
                    escaped = False
                    continue
                if char == "\\" and in_string:
                    escaped = True
                    continue
                if char == '"':
                    in_string = not in_string
                    continue
                if in_string:
                    continue
                if char in "{[":
                    balance += 1
                elif char in "}]":
                    balance -= 1
            if started and balance <= 0:
                break

        return "\n".join(collected) if collected else None

    def _find_first_json_char(self, line: str) -> int:
        positions = [pos for pos in (line.find("{"), line.find("[")) if pos >= 0]
        return min(positions) if positions else -1

    def _loads_json_with_comments(self, text: str) -> Any | None:
        cleaned = "\n".join(self._strip_json_line_comment(line) for line in text.splitlines())
        try:
            return json.loads(cleaned)
        except json.JSONDecodeError:
            return None

    def _strip_json_line_comment(self, line: str) -> str:
        in_string = False
        escaped = False
        for index, char in enumerate(line):
            if escaped:
                escaped = False
                continue
            if char == "\\" and in_string:
                escaped = True
                continue
            if char == '"':
                in_string = not in_string
                continue
            if not in_string and line[index : index + 2] == "--":
                return line[:index].rstrip()
        return line

    def _extract_extra_response_fields(self, segment_lines: list[str]) -> dict[str, dict[str, Any]]:
        fields: dict[str, dict[str, Any]] = {}
        for line in segment_lines:
            match = self.extra_response_field_pattern.search(line.strip())
            if not match:
                continue
            name = match.group(1)
            description = match.group(2).strip()
            field_type = "boolean" if "true" in description.lower() and "false" in description.lower() else "string"
            fields[name] = {
                "type": field_type,
                "description": description or None,
            }
        return fields

    def _build_openapi_responses(self, endpoint: dict[str, Any]) -> dict[str, Any]:
        schema: dict[str, Any] | None = None
        response_example = endpoint.get("response_example")
        if response_example is not None:
            schema = self._infer_schema(response_example)
        elif endpoint.get("extra_response_fields"):
            schema = {"type": "object", "properties": {}}

        if schema and endpoint.get("extra_response_fields"):
            properties = schema.setdefault("properties", {})
            for name, field_schema in endpoint["extra_response_fields"].items():
                properties[name] = {key: value for key, value in field_schema.items() if value is not None}

        response: dict[str, Any] = {"description": "成功"}
        if schema:
            media: dict[str, Any] = {"schema": schema}
            if response_example is not None:
                media["example"] = response_example
            response["content"] = {"application/json": media}
        return {"200": response}

    def _infer_schema(self, value: Any, required_for_object: bool = False) -> dict[str, Any]:
        if isinstance(value, bool):
            return {"type": "boolean"}
        if isinstance(value, int) and not isinstance(value, bool):
            return {"type": "integer"}
        if isinstance(value, float):
            return {"type": "number"}
        if isinstance(value, str):
            return {"type": "string"}
        if value is None:
            return {"nullable": True}
        if isinstance(value, list):
            item_schema = self._infer_schema(value[0]) if value else {}
            return {"type": "array", "items": item_schema}
        if isinstance(value, dict):
            schema: dict[str, Any] = {
                "type": "object",
                "properties": {key: self._infer_schema(item) for key, item in value.items()},
            }
            if required_for_object and value:
                schema["required"] = list(value.keys())
            return schema
        return {"type": "string"}

    def _endpoint_to_dict(self, endpoint: APIDocEndpointDraft) -> dict[str, Any]:
        return {
            "method": endpoint.method,
            "path": endpoint.path,
            "summary": endpoint.summary,
            "tag": endpoint.tag,
            "description": endpoint.description,
            "request_example": endpoint.request_example,
            "response_example": endpoint.response_example,
            "extra_response_fields": endpoint.extra_response_fields,
        }

    def _decode_text(self, content: bytes) -> str:
        for encoding in ("utf-8-sig", "utf-8", "gb18030"):
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        return content.decode("utf-8", errors="ignore")
